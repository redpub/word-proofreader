from __future__ import annotations

from dataclasses import dataclass
from typing import List
from io import BytesIO

from docx import Document
from docx_revisions import RevisionDocument, RevisionParagraph


@dataclass
class ResolvedChange:
    change_idx: int
    type: str
    author: str
    timestamp: str | None
    text: str
    start: int | None
    end: int | None
    position: int | None


def apply_changes_to_docx(
    doc_bytes: bytes,
    resolved_changes: List[ResolvedChange],
) -> bytes:
    """
    FIXED: Correct RevisionDocument workflow. No Streamlit imports.
    """
    input_stream = BytesIO(doc_bytes)
    rdoc = RevisionDocument(input_stream)
    
    # Build paragraph mapping (global char -> para_idx, local_offset)
    global_pos = 0
    para_mapping = {}
    
    for para_idx, para in enumerate(rdoc.paragraphs):
        para_text = para.text
        para_len = len(para_text)
        for local_offset in range(para_len + 1):  # +1 to allow insert at end
            para_mapping[global_pos + local_offset] = (para_idx, local_offset)
        global_pos += para_len + 1  # '\n' separator
    
    # Apply in reverse order
    sorted_changes = sorted(resolved_changes, key=lambda ch: -(ch.position or ch.start or 0))
    
    for ch in sorted_changes:
        try:
            if ch.type == "insert" and ch.position is not None and ch.position in para_mapping:
                para_idx, local_pos = para_mapping[ch.position]
                para = rdoc.paragraphs[para_idx]
                rp = RevisionParagraph.from_paragraph(para)
                rp.add_tracked_insertion(
                    text=ch.text,
                    author=ch.author,
                    revision_id=ch.change_idx,
                )
            
            elif ch.type in ("delete", "replace") and ch.start is not None and ch.end is not None:
                if ch.start in para_mapping:
                    para_idx, local_start = para_mapping[ch.start]
                    local_end = local_start + (ch.end - ch.start)
                    
                    para = rdoc.paragraphs[para_idx]
                    rp = RevisionParagraph.from_paragraph(para)
                    
                    if ch.type == "delete":
                        rp.add_tracked_deletion(
                            start=local_start,
                            end=local_end,
                            author=ch.author,
                            revision_id=ch.change_idx,
                        )
                    else:  # replace
                        rp.replace_tracked_at(
                            start=local_start,
                            end=local_end,
                            replace_text=ch.text,
                            author=ch.author,
                            revision_id=ch.change_idx,
                        )
        except Exception:
            continue  # Skip problematic changes
    
    # CRITICAL: Proper BytesIO handling
    output = BytesIO()
    rdoc.save(output)
    output.seek(0)
    result = output.read()
    
    return result
