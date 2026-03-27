import streamlit as st
import tempfile
import os
import json
import difflib
import time
import re
import traceback
import hashlib
import requests
from collections import OrderedDict
from typing import Any, List, Dict, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from openai import OpenAI
from docx_revisions import RevisionDocument, RevisionParagraph
from pydantic import BaseModel
from supabase import create_client, Client
from config import (
    POPULAR_MODELS,
    DEFAULT_CHUNK_SIZE,
    MAX_CHUNK_SIZE,
    DEFAULT_MAX_WORKERS,
    DEFAULT_MAX_RETRIES,
    DEFAULT_RETRY_DELAY,
    TOKEN_CUSHION_FACTOR,
    TOKENS_PER_CHAR_ESTIMATE,
    PROMPT_TEMPLATE_OVERHEAD,
    MIN_COMPLETION_RESERVE,
    DEFAULT_PARAGRAPHS_PER_PAGE,
    DEFAULT_EDITS_PER_PAGE
)

st.set_page_config(
    page_title="AI Word Proofreader",
    page_icon="📝",
    layout="wide"
)

class Edit(BaseModel):
    paragraph_index: int
    corrected_text: str
    reason: str

def compute_paragraph_hash(text: str) -> str:
    """Compute a short hash of paragraph text for verification."""
    return hashlib.md5(text.encode('utf-8')).hexdigest()[:8]

class ProofreadingResponse(BaseModel):
    edits: List[Edit]
    summary: str

# ============================================================================
# Prompt Management Functions (Supabase)
# ============================================================================

@st.cache_resource
def get_supabase_client() -> Client:
    """Initialize and cache the Supabase client."""
    url = st.secrets["connections"]["supabase"]["SUPABASE_URL"]
    key = st.secrets["connections"]["supabase"]["SUPABASE_KEY"]
    return create_client(url, key)

def load_prompts() -> Dict[str, Dict[str, Any]]:
    """
    Load prompts from Supabase.
    Returns ordered dict with default prompt first.
    """
    try:
        supabase = get_supabase_client()
        response = supabase.table("prompts").select("name, content, is_protected").order("created_at").execute()
        
        prompts = OrderedDict()
        default_name = "預設"
        
        # Put default prompt first
        for row in response.data:
            if row["name"] == default_name:
                prompts[default_name] = {
                    "content": row["content"],
                    "protected": row["is_protected"],
                }
                break
        
        # Add remaining prompts
        for row in response.data:
            if row["name"] != default_name:
                prompts[row["name"]] = {
                    "content": row["content"],
                    "protected": row["is_protected"],
                }
        
        if not prompts:
            st.error("無法載入提示：資料庫中無提示資料")
            st.stop()
        
        return prompts
    except Exception as e:
        st.error(f"無法載入提示：{str(e)}")
        st.stop()

_MAX_PROMPT_NAME_LENGTH = 50

def _validate_prompt_name(name: str) -> Tuple[bool, str]:
    """Validate prompt name."""
    name = name.strip()
    if not name:
        return False, "提示名稱不能為空"
    if len(name) > _MAX_PROMPT_NAME_LENGTH:
        return False, f"提示名稱不能超過 {_MAX_PROMPT_NAME_LENGTH} 個字元"
    return True, ""

def add_prompt(prompts: Dict[str, Dict[str, Any]], name: str, content: str) -> Tuple[bool, str]:
    """
    Add a new prompt to Supabase.
    Returns (success, message).
    """
    valid, msg = _validate_prompt_name(name)
    if not valid:
        return False, msg
    
    if not content or not content.strip():
        return False, "提示內容不能為空"
    
    if name in prompts:
        return False, f"提示 '{name}' 已存在"
    
    try:
        supabase = get_supabase_client()
        supabase.table("prompts").insert({
            "name": name,
            "content": content,
            "is_protected": False,
        }).execute()
        return True, f"已新增提示 '{name}'"
    except Exception as e:
        return False, f"儲存失敗：{str(e)}"

def update_prompt(prompts: Dict[str, Dict[str, Any]], name: str, content: str) -> Tuple[bool, str]:
    """
    Update an existing prompt's content in Supabase (if not protected).
    Returns (success, message).
    """
    if name not in prompts:
        return False, f"提示 '{name}' 不存在"
    
    if prompts[name].get("protected", False):
        return False, f"提示 '{name}' 受保護，無法編輯"
    
    if not content or not content.strip():
        return False, "提示內容不能為空"
    
    try:
        supabase = get_supabase_client()
        supabase.table("prompts").update({
            "content": content,
        }).eq("name", name).execute()
        return True, f"已更新提示 '{name}'"
    except Exception as e:
        return False, f"儲存失敗：{str(e)}"

def delete_prompt(prompts: Dict[str, Dict[str, Any]], name: str) -> Tuple[bool, str]:
    """
    Delete a prompt from Supabase (if not protected).
    Returns (success, message).
    """
    if name not in prompts:
        return False, f"提示 '{name}' 不存在"
    
    if prompts[name].get("protected", False):
        return False, f"提示 '{name}' 受保護，無法刪除"
    
    try:
        supabase = get_supabase_client()
        supabase.table("prompts").delete().eq("name", name).execute()
        return True, f"已刪除提示 '{name}'"
    except Exception as e:
        return False, f"刪除失敗：{str(e)}"

def get_openrouter_client(api_key: str) -> OpenAI:
    return OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key,
    )

def check_api_credits(api_key: str) -> Optional[Dict]:
    """
    Check OpenRouter API key credits.
    Returns dict with credit info or None if failed.
    """
    try:
        response = requests.get(
            "https://openrouter.ai/api/v1/auth/key",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=5
        )
        if response.status_code == 200:
            return response.json()
        return None
    except Exception:
        return None

def fetch_model_info(api_key: str, model_id: str) -> Optional[Dict]:
    """
    Fetch model metadata from OpenRouter API.
    Returns dict with context_length, max_completion_tokens, etc. or None if failed.
    """
    try:
        response = requests.get(
            "https://openrouter.ai/api/v1/models",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=10
        )
        if response.status_code != 200:
            return None
        
        data = response.json()
        for model in data.get("data", []):
            if model.get("id") == model_id:
                return {
                    "id": model["id"],
                    "name": model.get("name", model_id),
                    "context_length": model.get("context_length", 0),
                    "max_completion_tokens": model.get("top_provider", {}).get("max_completion_tokens", 0),
                    "tokenizer": model.get("architecture", {}).get("tokenizer", "unknown"),
                }
        return None
    except Exception:
        return None

def estimate_tokens(text: str) -> int:
    """
    Estimate token count for a string using a conservative CJK-aware heuristic.
    CJK characters typically map to 1-2 tokens; ASCII words ~1 token per 4 chars.
    We use a single conservative multiplier for safety.
    """
    return int(len(text) * TOKENS_PER_CHAR_ESTIMATE)

def compute_dynamic_chunk_size(
    rdoc: RevisionDocument,
    system_prompt: str,
    context_length: int,
    max_completion_tokens: int,
    process_percentage: int = 100
) -> int:
    """
    Compute optimal paragraphs-per-chunk based on model token limits.
    
    Budget: context_length - completion_reserve - system_prompt_tokens - prompt_overhead
    Then apply cushion factor and divide by average tokens per paragraph.
    Falls back to DEFAULT_CHUNK_SIZE if computation yields unreasonable results.
    """
    # Reserve tokens for the LLM's response
    completion_reserve = max(max_completion_tokens, MIN_COMPLETION_RESERVE)
    
    # Estimate system prompt tokens
    system_prompt_tokens = estimate_tokens(system_prompt)
    
    # Available tokens for paragraph content in the user prompt
    available = context_length - completion_reserve - system_prompt_tokens - PROMPT_TEMPLATE_OVERHEAD
    safe_budget = int(available * TOKEN_CUSHION_FACTOR)
    
    if safe_budget <= 0:
        return DEFAULT_CHUNK_SIZE
    
    # Estimate average tokens per paragraph from the actual document
    total_paragraphs = len(rdoc.paragraphs)
    paragraphs_to_process = max(1, int(total_paragraphs * process_percentage / 100))
    
    # Sample up to 200 paragraphs to get average size (avoid scanning huge documents fully)
    sample_size = min(paragraphs_to_process, 200)
    total_chars = 0
    for i in range(sample_size):
        text = rdoc.paragraphs[i].text.strip()
        if text:
            # Account for the "[index|hash] " prefix (~15 chars overhead per line)
            total_chars += len(text) + 15
        else:
            # Empty paragraph marker
            total_chars += 25
    
    if total_chars == 0:
        return DEFAULT_CHUNK_SIZE
    
    avg_tokens_per_para = int((total_chars / sample_size) * TOKENS_PER_CHAR_ESTIMATE)
    
    if avg_tokens_per_para <= 0:
        return DEFAULT_CHUNK_SIZE
    
    chunk_size = int(safe_budget / avg_tokens_per_para)
    
    # Clamp to reasonable bounds
    chunk_size = max(10, min(chunk_size, MAX_CHUNK_SIZE))
    
    return chunk_size

def check_for_tracked_changes(rdoc: RevisionDocument) -> Tuple[bool, int]:
    """
    Check if the document has any pending tracked changes (insertions or deletions).
    Returns (has_changes, count) where has_changes is True if tracked changes exist.
    """
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.ns import qn
    
    change_count = 0
    
    # Access the underlying document
    doc = rdoc._document
    
    # Check all paragraphs for tracked changes
    for paragraph in doc.paragraphs:
        # Check for insertions (w:ins)
        insertions = paragraph._element.findall(qn('w:ins'))
        change_count += len(insertions)
        
        # Check for deletions (w:del)
        deletions = paragraph._element.findall(qn('w:del'))
        change_count += len(deletions)
        
        # Check for move from/to
        move_from = paragraph._element.findall(qn('w:moveFrom'))
        change_count += len(move_from)
        
        move_to = paragraph._element.findall(qn('w:moveTo'))
        change_count += len(move_to)
    
    # Also check in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    insertions = paragraph._element.findall(qn('w:ins'))
                    change_count += len(insertions)
                    
                    deletions = paragraph._element.findall(qn('w:del'))
                    change_count += len(deletions)
                    
                    move_from = paragraph._element.findall(qn('w:moveFrom'))
                    change_count += len(move_from)
                    
                    move_to = paragraph._element.findall(qn('w:moveTo'))
                    change_count += len(move_to)
    
    return (change_count > 0, change_count)

def read_document_paragraphs(rdoc: RevisionDocument) -> str:
    lines = []
    for i, para in enumerate(rdoc.paragraphs):
        text = para.text.strip()
        if text:
            lines.append(f"[{i}] {text}")
    return "\n".join(lines)

def proofread_chunk_with_retry(
    client: OpenAI,
    model: str,
    chunk_text: str,
    system_prompt: str,
    chunk_info: str = "",
    max_retries: int = 3,
    initial_delay: float = 1.0,
    max_completion_tokens: Optional[int] = None
) -> Tuple[Optional[ProofreadingResponse], List[str]]:
    """
    Wrapper function that retries chunk processing with exponential backoff.
    Returns (result, warnings) tuple. Warnings are collected instead of calling
    st.warning() directly, since this function may run in worker threads where
    Streamlit calls are not thread-safe.
    """
    warnings = []
    for attempt in range(max_retries):
        try:
            result, chunk_warnings = proofread_chunk_with_llm(client, model, chunk_text, system_prompt, chunk_info, max_completion_tokens=max_completion_tokens)
            warnings.extend(chunk_warnings)
            if result is not None:
                return result, warnings
            
            # If result is None but no exception, still retry
            if attempt < max_retries - 1:
                delay = initial_delay * (2 ** attempt)
                warnings.append(f"重試 {attempt + 1}/{max_retries}{chunk_info}：未取得結果（等待 {delay} 秒）")
                time.sleep(delay)
                
        except Exception as e:
            if attempt < max_retries - 1:
                delay = initial_delay * (2 ** attempt)
                warnings.append(f"重試 {attempt + 1}/{max_retries}{chunk_info} 發生錯誤：{str(e)[:100]}...（等待 {delay} 秒）")
                time.sleep(delay)
            else:
                warnings.append(f"失敗{chunk_info}，已重試 {max_retries} 次：{str(e)}")
                return None, warnings
    
    return None, warnings

def chunk_paragraphs(rdoc: RevisionDocument, chunk_size: int = 100, process_percentage: int = 100) -> List[Tuple[int, int, str]]:
    """
    Split document into chunks for processing.
    Returns list of (start_index, end_index, text) tuples.
    
    Args:
        rdoc: RevisionDocument to process
        chunk_size: Number of paragraphs per chunk
        process_percentage: Percentage of document to process (1-100)
    
    IMPORTANT: Includes ALL paragraphs (even empty ones) to maintain correct indexing.
    Empty paragraphs are marked as [i] (empty) so the LLM knows to skip them.
    """
    chunks = []
    total_paragraphs = len(rdoc.paragraphs)
    
    # Calculate how many paragraphs to process based on percentage
    paragraphs_to_process = max(1, int(total_paragraphs * process_percentage / 100))
    
    for start_idx in range(0, paragraphs_to_process, chunk_size):
        end_idx = min(start_idx + chunk_size, paragraphs_to_process)
        
        lines = []
        for i in range(start_idx, end_idx):
            text = rdoc.paragraphs[i].text.strip()
            if text:
                h = compute_paragraph_hash(text)
                lines.append(f"[{i}|{h}] {text}")
            else:
                # Include empty paragraphs to maintain correct indexing
                lines.append(f"[{i}] (empty paragraph)")
        
        chunk_text = "\n".join(lines)
        chunks.append((start_idx, end_idx, chunk_text))
    
    return chunks

def proofread_chunk_with_llm(
    client: OpenAI,
    model: str,
    chunk_text: str,
    system_prompt: str,
    chunk_info: str = "",
    max_completion_tokens: Optional[int] = None
) -> Tuple[Optional[ProofreadingResponse], List[str]]:
    """
    Process a single chunk with the LLM.
    Returns (result, warnings) tuple. Warnings are collected instead of calling
    st.*() directly, since this function may run in worker threads where
    Streamlit calls are not thread-safe.
    """
    warnings = []
    try:
        user_prompt = f"""Here is the document chunk to proofread{chunk_info}:

{chunk_text}

Please analyze this document and provide corrections in the following JSON format.
IMPORTANT: Your response must be VALID JSON only, with no additional text before or after.

{{
    "edits": [
        {{
            "paragraph_index": 0,
            "corrected_text": "the complete corrected paragraph text",
            "reason": "why this change improves the text"
        }}
    ],
    "summary": "所有修改的簡短摘要（繁體中文）"
}}

RULES:
1. Each paragraph is shown as [index|hash] text. Return the paragraph_index (the number before the |) for each edit.
2. Provide the COMPLETE corrected paragraph text in corrected_text (the full paragraph with your corrections applied).
3. Only include paragraphs that need corrections.
4. SKIP paragraphs marked as "(empty paragraph)" - do not suggest edits for them.
5. The corrected_text should contain ONLY the corrected text, NOT the [index|hash] prefix.
6. Ensure all strings are properly escaped for JSON (escape quotes, newlines, etc.)
7. If no corrections are needed, return: {{"edits": [], "summary": "無需修正"}}
8. The "summary" and "reason" fields MUST be written in Traditional Chinese (繁體中文).

EXAMPLE:
Input: "[102|a3f9b2c1] 對於我來説，語文很難。"
Correct JSON:
{{
    "paragraph_index": 102,
    "corrected_text": "對於我來說，語文很難。",
    "reason": "將「説」修正為「說」"
}}

Respond with valid JSON only."""

        create_kwargs = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.3,
            "response_format": {"type": "json_object"}
        }
        if max_completion_tokens:
            create_kwargs["max_tokens"] = max_completion_tokens
        
        response = client.chat.completions.create(**create_kwargs)
        
        content = response.choices[0].message.content
        
        if not content:
            # Gather diagnostic info from the response
            choice = response.choices[0] if response.choices else None
            diag = {
                "finish_reason": getattr(choice, "finish_reason", None) if choice else None,
                "refusal": getattr(choice.message, "refusal", None) if choice else None,
                "model": getattr(response, "model", None),
                "usage": getattr(response, "usage", None),
            }
            warnings.append(f"LLM 回傳空白回應{chunk_info}。診斷資訊：{diag}")
            return None, warnings
        
        # Check for truncation before parsing
        finish_reason = response.choices[0].finish_reason
        if finish_reason == "length":
            warnings.append(f"LLM 回應被截斷（finish_reason=length）{chunk_info}。"
                            f"區塊可能超出模型的輸出上限。")
            return None, warnings
        
        try:
            data = json.loads(content)
        except json.JSONDecodeError as je:
            warnings.append(f"無法解析 LLM 回應的 JSON{chunk_info}：{str(je)}")
            return None, warnings
        
        if "edits" not in data:
            warnings.append(f"LLM 回應缺少 'edits' 欄位{chunk_info}")
            return ProofreadingResponse(edits=[], summary=data.get("summary", "未提供修改")), warnings
        
        try:
            edits = [Edit(**edit) for edit in data.get("edits", [])]
        except Exception as validation_error:
            warnings.append(f"驗證修改資料失敗{chunk_info}：{str(validation_error)}")
            return None, warnings
        
        return ProofreadingResponse(
            edits=edits,
            summary=data.get("summary", "無需修正。")
        ), warnings
    except Exception as e:
        warnings.append(f"呼叫 LLM 時發生錯誤{chunk_info}：{str(e)}")
        return None, warnings

def proofread_with_llm(
    client: OpenAI,
    model: str,
    rdoc: RevisionDocument,
    system_prompt: str,
    max_workers: int = 5,
    process_percentage: int = 100,
    model_info: Optional[Dict] = None
) -> Optional[ProofreadingResponse]:
    """
    Proofread document in chunks using parallel processing.
    Chunk size is computed dynamically from model token limits when model_info is available.
    """
    total_paragraphs = len(rdoc.paragraphs)
    paragraphs_to_process = max(1, int(total_paragraphs * process_percentage / 100))
    
    # Compute chunk size dynamically or fall back to default
    if model_info and model_info.get("context_length", 0) > 0:
        chunk_size = compute_dynamic_chunk_size(
            rdoc, system_prompt,
            model_info["context_length"],
            model_info.get("max_completion_tokens", 0),
            process_percentage
        )
        st.info(f"📐 動態區塊大小：{chunk_size} 個段落"
                f"（模型上下文：{model_info['context_length']:,} tokens）")
    else:
        chunk_size = DEFAULT_CHUNK_SIZE
        st.info(f"📐 使用預設區塊大小：{chunk_size} 個段落（無法取得模型資訊）")
    
    chunks = chunk_paragraphs(rdoc, chunk_size, process_percentage)
    total_chunks = len(chunks)
    
    if total_chunks == 0:
        return ProofreadingResponse(edits=[], summary="文件為空")
    
    if process_percentage < 100:
        st.info(f"🧪 測試模式：處理文件的 {process_percentage}%（{paragraphs_to_process}/{total_paragraphs} 個段落）")
    
    actual_workers = min(max_workers, total_chunks)
    st.info(f"📦 將文件分為 {total_chunks} 個區塊（每區塊 {chunk_size} 個段落），使用 {actual_workers} 個平行工作執行緒")
    
    all_edits = []
    chunk_summaries = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    completed_chunks = 0
    
    # Process chunks in parallel with retry logic
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all tasks
        future_to_chunk = {}
        for chunk_idx, (start_idx, end_idx, chunk_text) in enumerate(chunks):
            chunk_info = f" (chunk {chunk_idx + 1}/{total_chunks}, paragraphs {start_idx}-{end_idx})"
            future = executor.submit(
                proofread_chunk_with_retry,
                client,
                model,
                chunk_text,
                system_prompt,
                chunk_info,
                max_retries=DEFAULT_MAX_RETRIES,
                initial_delay=DEFAULT_RETRY_DELAY,
                max_completion_tokens=model_info.get("max_completion_tokens") if model_info else None
            )
            future_to_chunk[future] = (chunk_idx, start_idx, end_idx)
        
        # Collect results as they complete
        chunk_results = {}
        all_warnings = []
        for future in as_completed(future_to_chunk):
            chunk_idx, start_idx, end_idx = future_to_chunk[future]
            completed_chunks += 1
            
            status_text.text(f"已完成 {completed_chunks}/{total_chunks} 個區塊（最新：段落 {start_idx}-{end_idx}）...")
            
            try:
                result, warnings = future.result()
                all_warnings.extend(warnings)
                if result:
                    chunk_results[chunk_idx] = result
            except Exception as e:
                all_warnings.append(f"處理區塊 {chunk_idx + 1} 時發生錯誤：{str(e)}")
            
            progress_bar.progress(completed_chunks / total_chunks)
    
    # Display collected warnings from the main thread (thread-safe)
    for warning_msg in all_warnings:
        st.warning(warning_msg)
    
    # Combine results in order
    for chunk_idx in sorted(chunk_results.keys()):
        result = chunk_results[chunk_idx]
        all_edits.extend(result.edits)
        if result.summary and result.summary != "無需修正。":
            chunk_summaries.append(f"區塊 {chunk_idx + 1}：{result.summary}")
    
    status_text.text("✅ 所有區塊處理完成！")
    
    # Combine summaries
    if chunk_summaries:
        combined_summary = f"已平行處理 {total_chunks} 個區塊。\n" + "\n".join(chunk_summaries)
    else:
        combined_summary = f"已平行處理 {total_chunks} 個區塊，無需修正。"
    
    return ProofreadingResponse(
        edits=all_edits,
        summary=combined_summary
    )

def compute_character_diffs(original: str, corrected: str) -> List[Tuple[str, int, int, str]]:
    """
    Compute character-level differences between original and corrected text.
    Returns list of (operation, start, end, text) tuples:
    - ('delete', start, end, deleted_text)
    - ('insert', position, position, inserted_text)
    - ('replace', start, end, new_text)
    """
    diffs = []
    matcher = difflib.SequenceMatcher(None, original, corrected)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'delete':
            diffs.append(('delete', i1, i2, original[i1:i2]))
        elif tag == 'insert':
            diffs.append(('insert', i1, i1, corrected[j1:j2]))
        elif tag == 'replace':
            diffs.append(('replace', i1, i2, corrected[j1:j2]))
    
    return diffs

def apply_tracked_changes(
    rdoc: RevisionDocument,
    edits: List[Edit],
    author: str
) -> Dict[str, Any]:
    stats = {
        "deletions": 0, 
        "insertions": 0, 
        "errors": 0,
        "failed_edits": []
    }
    
    for edit in edits:
        try:
            if edit.paragraph_index < 0 or edit.paragraph_index >= len(rdoc.paragraphs):
                stats["errors"] += 1
                stats["failed_edits"].append({
                    "paragraph_index": edit.paragraph_index,
                    "reason": "out_of_range",
                    "total_paragraphs": len(rdoc.paragraphs),
                    "edit": edit.model_dump()
                })
                continue
            
            para_element = rdoc.paragraphs[edit.paragraph_index]
            current_text = para_element.text
            
            # Use the actual document text as the original for diffing
            original_text = current_text
            
            rp = RevisionParagraph.from_paragraph(para_element)
            diffs = compute_character_diffs(original_text, edit.corrected_text)
            
            # Process diffs in reverse order to maintain correct positions
            for op, start, end, text in reversed(diffs):
                if op == 'delete':
                    rp.add_tracked_deletion(
                        start=start,
                        end=end,
                        author=author
                    )
                    stats["deletions"] += 1
                    
                elif op == 'insert':
                    # For pure insertions, we need to insert at a position.
                    # replace_tracked_at requires start < end, so we "borrow"
                    # an adjacent character, delete it, and re-insert it
                    # alongside the new text.
                    if start >= len(original_text) or len(original_text) == 0:
                        # Insert at end or into empty paragraph
                        rp.add_tracked_insertion(
                            text=text,
                            author=author
                        )
                    elif start == 0:
                        # Insert at beginning - borrow the first character
                        rp.replace_tracked_at(
                            start=0,
                            end=1,
                            replace_text=text + original_text[0],
                            author=author
                        )
                    else:
                        # Insert in middle - borrow the character at position
                        rp.replace_tracked_at(
                            start=start,
                            end=start + 1,
                            replace_text=text + original_text[start],
                            author=author
                        )
                    stats["insertions"] += 1
                    
                elif op == 'replace':
                    # Use replace_tracked_at which handles both deletion and insertion at position
                    rp.replace_tracked_at(
                        start=start,
                        end=end,
                        replace_text=text,
                        author=author
                    )
                    stats["deletions"] += 1
                    stats["insertions"] += 1
                    
        except Exception as e:
            stats["errors"] += 1
            stats["failed_edits"].append({
                "paragraph_index": edit.paragraph_index,
                "reason": "exception",
                "error_message": str(e),
                "corrected_text": edit.corrected_text,
                "edit_reason": edit.reason
            })
    
    return stats

def main():
    st.title("📝 紅出版 Word AI校對工具")
    st.markdown("上傳 Word 文件，讓 AI 進行校對並以追蹤修訂模式標記修改")
    
    # Load prompts
    if 'prompts' not in st.session_state:
        st.session_state.prompts = load_prompts()
    
    prompts = st.session_state.prompts
    prompt_names = list(prompts.keys())
    
    with st.sidebar:
        st.header("⚙️ 設定")
        
        api_key = st.text_input(
            "OpenRouter API 金鑰",
            type="password",
            help="從 https://openrouter.ai/keys 取得您的 API 金鑰"
        )
        
        # Show API credits if key is entered
        if api_key:
            # Only re-check credits when API key changes or on first load
            need_check = False
            if 'cached_api_key' not in st.session_state or st.session_state.cached_api_key != api_key:
                need_check = True
            
            if need_check:
                with st.spinner("檢查 API 額度..."):
                    credits_info = check_api_credits(api_key)
                st.session_state.cached_api_key = api_key
                st.session_state.cached_credits_info = credits_info
            else:
                credits_info = st.session_state.get('cached_credits_info')
            
            if credits_info and 'data' in credits_info:
                data = credits_info['data']
                # Display credit balance
                if 'limit' in data and data['limit'] is not None:
                    limit = float(data.get('limit', 0))
                    usage = float(data.get('usage', 0))
                    remaining = limit - usage
                    
                    if remaining > 0:
                        st.success(f"💰 剩餘額度：${remaining:.2f}")
                    else:
                        st.error(f"⚠️ 額度已用完")
                    
                    # Show usage bar
                    if limit > 0:
                        usage_percent = (usage / limit) * 100
                        st.progress(min(usage_percent / 100, 1.0))
                        st.caption(f"已使用：${usage:.2f} / ${limit:.2f} ({usage_percent:.1f}%)")
                else:
                    # Unlimited or rate-limited key
                    st.info("✓ API 金鑰有效")
            elif credits_info is None:
                st.warning("⚠️ 無法檢查 API 額度")
        
        model = st.selectbox(
            "模型",
            options=POPULAR_MODELS,
            index=0,
            help="選擇用於校對的 LLM 模型"
        )
        
        author_name = st.text_input(
            "作者名稱",
            value="紅出版",
            help="顯示在追蹤修訂中的名稱"
        )
        
        # ============================================================
        # Prompt Management
        # ============================================================
        st.markdown("---")
        st.subheader("📋 校對提示")
        
        # Initialize mode state
        if 'creating_new_prompt' not in st.session_state:
            st.session_state.creating_new_prompt = False
        
        # Mode toggle button
        if not st.session_state.creating_new_prompt:
            if st.button("➕ 新增提示", use_container_width=True, type="secondary"):
                st.session_state.creating_new_prompt = True
                st.session_state.new_prompt_name_input = ""
                st.session_state.confirm_delete_target = None
                st.rerun()
        else:
            if st.button("⬅️ 返回", use_container_width=True, type="secondary"):
                st.session_state.creating_new_prompt = False
                st.session_state.confirm_delete_target = None
                st.rerun()
        
        # ============================================================
        # NEW PROMPT MODE
        # ============================================================
        if st.session_state.creating_new_prompt:
            st.info("🆕 新增模式：輸入名稱並編輯內容")
            
            # New prompt name input
            new_prompt_name = st.text_input(
                "新提示名稱",
                value=st.session_state.get('new_prompt_name_input', ''),
                key="new_prompt_name"
            )
            
            # Template selector - use existing prompts as templates
            template_options = ["(空白)"] + list(prompts.keys())
            
            def on_template_change():
                selected = st.session_state.template_selector
                if selected != "(空白)":
                    st.session_state.new_prompt_content = prompts[selected]["content"]
            
            selected_template = st.selectbox(
                "從範本開始",
                options=template_options,
                key="template_selector",
                on_change=on_template_change
            )
            
            # Reuse the textarea UI for new prompt content
            new_prompt_content = st.text_area(
                "提示內容",
                value="",
                height=250,
                help="編輯新提示的內容",
                key="new_prompt_content"
            )
            
            # Create button
            if st.button("💾 建立提示", use_container_width=True, type="secondary", key="create_new"):
                if not new_prompt_name or not new_prompt_name.strip():
                    st.error("請輸入提示名稱")
                elif not new_prompt_content or not new_prompt_content.strip():
                    st.error("請輸入提示內容")
                else:
                    success, message = add_prompt(prompts, new_prompt_name, new_prompt_content)
                    if success:
                        st.success(message)
                        st.session_state.prompts = load_prompts()
                        st.session_state.creating_new_prompt = False
                        # Store the newly created prompt name to auto-select it
                        st.session_state.newly_created_prompt = new_prompt_name
                        st.session_state.confirm_delete_target = None
                        st.rerun()
                    else:
                        st.error(message)
            
            # Use empty content for system prompt in new mode
            system_prompt = new_prompt_content
        
        # ============================================================
        # EDIT EXISTING PROMPT MODE
        # ============================================================
        else:
            # If we need to change the selectbox value, do it BEFORE the widget renders
            if 'newly_created_prompt' in st.session_state:
                st.session_state.prompt_selector = st.session_state.newly_created_prompt
                del st.session_state.newly_created_prompt
            
            if 'reset_prompt_selector' in st.session_state:
                st.session_state.prompt_selector = st.session_state.reset_prompt_selector
                del st.session_state.reset_prompt_selector
            
            # Callback when user changes selection - clear delete confirmation
            def on_selector_change():
                st.session_state.confirm_delete_target = None
            
            # Prompt selector - controlled via key "prompt_selector"
            selected_prompt_name = st.selectbox(
                "選擇提示",
                options=prompt_names,
                key="prompt_selector",
                on_change=on_selector_change,
                help="選擇要使用的提示範本"
            )
            
            # Get selected prompt details
            selected_prompt_data = prompts[selected_prompt_name]
            is_protected = selected_prompt_data.get("protected", False)
            original_content = selected_prompt_data["content"]
            
            # Callback to track content changes
            def on_prompt_change():
                st.session_state.prompt_modified = True
            
            # Prompt content textarea (read-only for protected prompts)
            current_prompt_content = st.text_area(
                "提示內容",
                value=original_content,
                height=250,
                help="此提示受保護，無法編輯" if is_protected else "編輯提示內容",
                key=f"prompt_content_{selected_prompt_name}",
                disabled=is_protected,
                on_change=on_prompt_change
            )
            
            if is_protected:
                st.caption("🔒 此提示受保護，無法儲存修改或刪除")
            
            # Check if content has been modified
            content_modified = current_prompt_content != original_content
            
            # Save button - only enabled when content is modified and not protected
            save_disabled = is_protected or not content_modified
            if st.button(
                f"💾 儲存",
                disabled=save_disabled,
                use_container_width=True,
                key="save_current",
                type="primary" if content_modified and not is_protected else "secondary"
            ):
                success, message = update_prompt(prompts, selected_prompt_name, current_prompt_content)
                if success:
                    st.success(message)
                    st.session_state.prompts = load_prompts()
                    st.session_state.prompt_modified = False
                    st.rerun()
                else:
                    st.error(message)
            
            # Delete button with confirmation
            delete_disabled = is_protected
            
            # Initialize confirmation state
            if 'confirm_delete_target' not in st.session_state:
                st.session_state.confirm_delete_target = None
            
            # Check if we're confirming deletion for the currently selected prompt
            is_confirming = st.session_state.confirm_delete_target == selected_prompt_name
            
            if not is_confirming:
                if st.button(
                    f"🗑️ 刪除",
                    disabled=delete_disabled,
                    use_container_width=True,
                    key="delete_current"
                ):
                    st.session_state.confirm_delete_target = selected_prompt_name
                    st.rerun()
            else:
                st.warning(f"⚠️ 確定要刪除 '{selected_prompt_name}' 嗎？")
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("✅ 確認刪除", use_container_width=True, type="primary", key="confirm_delete_yes"):
                        success, message = delete_prompt(prompts, selected_prompt_name)
                        if success:
                            st.success(message)
                            st.session_state.prompts = load_prompts()
                            st.session_state.confirm_delete_target = None
                            # Schedule selector reset for next rerun (can't modify after widget renders)
                            st.session_state.reset_prompt_selector = prompt_names[0]
                            st.rerun()
                        else:
                            st.error(message)
                            st.session_state.confirm_delete_target = None
                
                with col2:
                    if st.button("❌ 取消", use_container_width=True, key="confirm_delete_no"):
                        st.session_state.confirm_delete_target = None
                        st.rerun()
            
            # Set the system prompt to use (use current edited content)
            system_prompt = current_prompt_content
        
        st.markdown("---")
        
        with st.expander("🧪 測試選項", expanded=False):
            process_percentage = st.slider(
                "處理文件的百分比",
                min_value=1,
                max_value=100,
                value=100,
                step=1,
                help="測試用：僅處理文件的一部分（例如 10% = 前 10% 的段落）"
            )
            st.caption(f"將處理文件的 {process_percentage}%")
        
        st.markdown("---")
        st.markdown("### 關於")
        st.markdown("本應用程式使用 AI 校對 Word 文件，並以追蹤修訂模式加入修正。")        
    
    uploaded_file = st.file_uploader(
        "上傳 Word 文件 (.docx)",
        type=["docx"],
        help="選擇要校對的 .docx 檔案"
    )
    
    if uploaded_file is not None:
        if not api_key:
            st.warning("⚠️ 請在側邊欄輸入您的 OpenRouter API 金鑰")
            return
        
        tmp_input_path = None
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_input:
            tmp_input_path = tmp_input.name
            tmp_input.write(uploaded_file.read())
        
        try:
            with st.spinner("📖 讀取文件中..."):
                rdoc = RevisionDocument(tmp_input_path)
                document_text = read_document_paragraphs(rdoc)
            
            # Check for pending tracked changes
            has_changes, change_count = check_for_tracked_changes(rdoc)
            if has_changes:
                st.error(f"❌ 此文件包含 {change_count} 個未處理的追蹤修訂")
                st.warning("⚠️ 請先在 Microsoft Word 中接受或拒絕所有追蹤修訂，然後重新上傳文件。")
                st.info("💡 在 Word 中：審閱 → 接受 → 接受所有修訂（或逐一檢視）")
                return
            
            st.success(f"✅ 已載入文件，共 {len(rdoc.paragraphs)} 個段落")
            
            with st.expander("📄 文件預覽", expanded=False):
                show_full = st.checkbox("顯示完整文件", value=False)
                
                if show_full:
                    paragraphs_list = document_text.split('\n')
                    total_paragraphs = len(paragraphs_list)
                    
                    # Pagination settings
                    paragraphs_per_page = DEFAULT_PARAGRAPHS_PER_PAGE
                    total_pages = (total_paragraphs + paragraphs_per_page - 1) // paragraphs_per_page
                    
                    page = st.number_input(
                        f"頁面 (1-{total_pages})",
                        min_value=1,
                        max_value=total_pages,
                        value=1,
                        step=1
                    )
                    
                    start_idx = (page - 1) * paragraphs_per_page
                    end_idx = min(start_idx + paragraphs_per_page, total_paragraphs)
                    
                    st.caption(f"顯示第 {start_idx + 1}-{end_idx} 行，共 {total_paragraphs} 行（非空段落）")
                    st.text('\n'.join(paragraphs_list[start_idx:end_idx]))
                else:
                    st.text(document_text[:2000] + ("..." if len(document_text) > 2000 else ""))
                    if len(document_text) > 2000:
                        st.caption(f"顯示前 2000 個字元。勾選「顯示完整文件」以查看更多內容。")
            
            if st.button("🚀 開始校對", type="primary", use_container_width=True):
                client = get_openrouter_client(api_key)
                
                # Fetch and cache model info for dynamic chunk sizing
                cache_key = f"model_info_{model}"
                if cache_key not in st.session_state:
                    with st.spinner("📡 取得模型資訊..."):
                        st.session_state[cache_key] = fetch_model_info(api_key, model)
                model_info = st.session_state[cache_key]
                
                with st.spinner(f"🤖 使用 {model} 校對中..."):
                    result = proofread_with_llm(
                        client,
                        model,
                        rdoc,
                        system_prompt,
                        max_workers=DEFAULT_MAX_WORKERS,
                        process_percentage=process_percentage,
                        model_info=model_info
                    )
                
                if result is None:
                    st.error("❌ 無法取得校對結果")
                    return
                
                if not result.edits:
                    st.info("✨ 無需修正！您的文件看起來很棒。")
                    st.markdown(f"**AI 摘要：** {result.summary}")
                    return
                
                st.success(f"✅ 找到 {len(result.edits)} 個建議修正")
                st.markdown(f"**摘要：** {result.summary}")
                
                with st.expander("📝 建議修改", expanded=True):
                    # Add pagination for edits if there are many
                    total_edits = len(result.edits)
                    
                    if total_edits > DEFAULT_EDITS_PER_PAGE:
                        edits_per_page = DEFAULT_EDITS_PER_PAGE
                        total_edit_pages = (total_edits + edits_per_page - 1) // edits_per_page
                        
                        edit_page = st.number_input(
                            f"修改頁面 (1-{total_edit_pages})",
                            min_value=1,
                            max_value=total_edit_pages,
                            value=1,
                            step=1,
                            key="edit_page"
                        )
                        
                        start_edit = (edit_page - 1) * edits_per_page
                        end_edit = min(start_edit + edits_per_page, total_edits)
                        edits_to_show = result.edits[start_edit:end_edit]
                        edit_offset = start_edit
                    else:
                        edits_to_show = result.edits
                        edit_offset = 0
                    
                    for i, edit in enumerate(edits_to_show, edit_offset + 1):
                        st.markdown(f"**修改 {i}** (段落 {edit.paragraph_index})")
                        
                        # Look up original text from document
                        if 0 <= edit.paragraph_index < len(rdoc.paragraphs):
                            original_text = rdoc.paragraphs[edit.paragraph_index].text
                        else:
                            original_text = "(段落索引超出範圍)"
                        
                        diffs = compute_character_diffs(original_text, edit.corrected_text)
                        
                        st.markdown("**原文：**")
                        st.code(original_text, language=None)
                        st.markdown("**修正後：**")
                        st.code(edit.corrected_text, language=None)
                        
                        st.markdown("**字元級別修改：**")
                        for op, start, end, text in diffs:
                            if op == 'delete':
                                st.markdown(f"🔴 刪除位置 {start}-{end}：`{text}`")
                            elif op == 'insert':
                                st.markdown(f"🟢 插入位置 {start}：`{text}`")
                            elif op == 'replace':
                                st.markdown(f"🟡 替換位置 {start}-{end} 為：`{text}`")
                        
                        st.caption(f"💡 {edit.reason}")
                        st.markdown("---")
                
                with st.spinner("✏️ 套用追蹤修訂中..."):
                    stats = apply_tracked_changes(rdoc, result.edits, author_name)
                
                st.info(f"📊 已套用 {stats['deletions']} 個刪除和 {stats['insertions']} 個插入")
                
                if stats['errors'] > 0:
                    st.warning(f"⚠️ {stats['errors']} 個修改無法套用")
                    
                    # Debug interface for failed edits
                    with st.expander("🔍 除錯：失敗修改詳情", expanded=False):
                        st.markdown("### 失敗修改分析")
                        st.caption(f"總失敗數：{len(stats['failed_edits'])}")
                        
                        # Group by failure reason
                        range_failures = [f for f in stats['failed_edits'] if f['reason'] == 'out_of_range']
                        exception_failures = [f for f in stats['failed_edits'] if f['reason'] == 'exception']
                        
                        st.markdown(f"- **超出範圍**：{len(range_failures)}")
                        st.markdown(f"- **例外錯誤**：{len(exception_failures)}")
                        
                        # Show out of range failures
                        if range_failures:
                            st.markdown("### 📍 超出範圍失敗")
                            for failure in range_failures:
                                st.markdown(f"- 段落 {failure['paragraph_index']}（文件共有 {failure['total_paragraphs']} 個段落）")
                        
                        # Show exception failures
                        if exception_failures:
                            st.markdown("### ⚠️ 例外錯誤失敗")
                            for i, failure in enumerate(exception_failures[:10], 1):
                                st.markdown(f"**例外 {i}** - 段落 {failure['paragraph_index']}")
                                st.code(failure['error_message'])
                                st.caption(f"修正內容：{failure['corrected_text'][:100]}")
                                st.markdown("---")
                        
                        # Export debug data
                        st.markdown("### 💾 匯出除錯資料")
                        debug_json = json.dumps(stats['failed_edits'], indent=2, ensure_ascii=False)
                        st.download_button(
                            label="下載失敗修改 JSON",
                            data=debug_json,
                            file_name="failed_edits_debug.json",
                            mime="application/json"
                        )
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_output:
                    tmp_output_path = tmp_output.name
                
                rdoc.save(tmp_output_path)
                
                with open(tmp_output_path, "rb") as f:
                    output_data = f.read()
                
                original_name = uploaded_file.name.rsplit(".", 1)[0]
                output_filename = f"{original_name}_proofread.docx"
                
                st.download_button(
                    label="⬇️ 下載校對後的文件",
                    data=output_data,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
                
                st.success("✅ 文件已準備好！在 Microsoft Word 中開啟以檢視追蹤修訂。")
                
                os.unlink(tmp_output_path)
        
        except Exception as e:
            st.error(f"❌ 處理文件時發生錯誤：{str(e)}")
        finally:
            if tmp_input_path and os.path.exists(tmp_input_path):
                os.unlink(tmp_input_path)
    else:        
        
        st.markdown("### 使用方法")
        st.markdown("""
        1. **上傳**您的 Word 文件 (.docx)
        2. **設定** AI 模型和校對提示
        3. **檢視**建議的修正
        4. **下載**帶有追蹤修訂的文件
        5. **在 Microsoft Word 中開啟**以接受/拒絕修改
        """)

if __name__ == "__main__":
    main()
